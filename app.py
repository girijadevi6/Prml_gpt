import streamlit as st
import sys
import os
from io import BytesIO
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# -------------------------------------------------
# Import backend
# -------------------------------------------------
sys.path.append(os.path.join(os.getcwd(), "src"))
from prml_qa import generate_answer, retrieve_chunks

# -------------------------------------------------
# Page config
# -------------------------------------------------
st.set_page_config(
    page_title="PRML QA Chat",
    page_icon="🧠",
    layout="wide"
)

st.title("🧠 PRML Question Answering Chat")
st.caption("Ask PRML questions and get clear, structured answers.")

# -------------------------------------------------
# Session state
# -------------------------------------------------
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

# -------------------------------------------------
# Sidebar
# -------------------------------------------------
st.sidebar.header("⚙️ Settings")
show_context = st.sidebar.checkbox("Show retrieved context", value=False)

if st.sidebar.button("🧹 Clear Chat"):
    st.session_state.chat_history = []
    st.rerun()

# -------------------------------------------------
# Display chat history
# -------------------------------------------------
for msg in st.session_state.chat_history:
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])

# -------------------------------------------------
# Chat input
# -------------------------------------------------
user_query = st.chat_input("Type your PRML question here...")

if user_query:
    # -------- Store user question --------
    user_block = f"""
### ❓ Question
{user_query}
"""
    st.session_state.chat_history.append(
        {"role": "user", "content": user_block}
    )

    with st.chat_message("assistant"):
        with st.spinner("Generating answer..."):
            answer = generate_answer(user_query)

        # -------- Neat formatted answer --------
        answer_block = f"""
### ✅ Answer
{answer}
"""
        st.markdown(answer_block)

        # Save assistant response
        st.session_state.chat_history.append(
            {"role": "assistant", "content": answer_block}
        )

        # -------------------------------------------------
        # Download buttons
        # -------------------------------------------------
        col1, col2 = st.columns(2)

        with col1:
            def word_bytes(text):
                doc = Document()
                doc.add_heading("PRML Question", level=1)
                doc.add_paragraph(user_query)
                doc.add_heading("Answer", level=1)
                doc.add_paragraph(text)
                bio = BytesIO()
                doc.save(bio)
                return bio.getvalue()

            st.download_button(
                "📄 Download Word",
                data=word_bytes(answer),
                file_name="prml_answer.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        with col2:
            def pdf_bytes(text):
                bio = BytesIO()
                c = canvas.Canvas(bio, pagesize=letter)
                width, height = letter
                y = height - 40

                c.setFont("Helvetica-Bold", 14)
                c.drawString(40, y, "PRML Question")
                y -= 20
                c.setFont("Helvetica", 11)
                c.drawString(40, y, user_query)
                y -= 30

                c.setFont("Helvetica-Bold", 14)
                c.drawString(40, y, "Answer")
                y -= 20
                c.setFont("Helvetica", 11)

                for line in text.split("\n"):
                    c.drawString(40, y, line)
                    y -= 14
                    if y < 40:
                        c.showPage()
                        y = height - 40

                c.save()
                bio.seek(0)
                return bio.read()

            st.download_button(
                "📄 Download PDF",
                data=pdf_bytes(answer),
                file_name="prml_answer.pdf",
                mime="application/pdf"
            )

        # -------------------------------------------------
        # Optional context
        # -------------------------------------------------
        if show_context:
            st.markdown("---")
            st.subheader("🔍 Retrieved Context")
            chunks = retrieve_chunks(user_query)
            for i, chunk in enumerate(chunks, 1):
                with st.expander(f"Chunk {i}"):
                    st.write(chunk["text"])
